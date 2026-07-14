import os
from datetime import datetime
from django.views.generic import TemplateView, View
from django.http import HttpResponse
from django.contrib.staticfiles.finders import find

from apps.base.models import TipoRU, CarreraPostgrado

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING


# ══════════════════════════════════════════════════════
# MÓDULO GENERADOR — Creación asistida de RU en formato word
# Templates: ru_crear.html
# ══════════════════════════════════════════════════════

class ResolucionCrearPaginaView(TemplateView):
    template_name = "generador/ru_crear.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["tipos"] = TipoRU.objects.all()
        context["carreras"] = CarreraPostgrado.objects.all()
        return context


class ResolucionUniversitariaGenerarWordView(View):
    def post(self, request, *args, **kwargs):
        titulo = request.POST.get('titulo', '').strip().upper()
        anio = request.POST.get('anio', '2026')
        numero_ru = request.POST.get('numero_ru', '____')
        vistos_textos = request.POST.getlist('vistos_textos')
        considerandos_extra = request.POST.getlist('considerandos')
        resuelvo_texto = request.POST.get('resuelvo', '').strip()

        doc = Document()

        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1)

        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)

        def set_run_font(run, bold=False):
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = bold

        def set_paragraph_compact(p):
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.line_spacing = 1.0

        def add_paragraph_border(p, top=False, bottom=False):
            p_pr = p._p.get_or_add_pPr()
            p_bdr = p_pr.find(qn('w:pBdr'))

            if p_bdr is None:
                p_bdr = OxmlElement('w:pBdr')
                p_pr.append(p_bdr)

            if top:
                top_bdr = OxmlElement('w:top')
                top_bdr.set(qn('w:val'), 'single')
                top_bdr.set(qn('w:sz'), '6')
                top_bdr.set(qn('w:space'), '1')
                top_bdr.set(qn('w:color'), '000000')
                p_bdr.append(top_bdr)

            if bottom:
                bottom_bdr = OxmlElement('w:bottom')
                bottom_bdr.set(qn('w:val'), 'single')
                bottom_bdr.set(qn('w:sz'), '6')
                bottom_bdr.set(qn('w:space'), '1')
                bottom_bdr.set(qn('w:color'), '000000')
                p_bdr.append(bottom_bdr)
                
        def add_section_title(text):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(text)
            set_run_font(run, bold=True)

        meses = [
            "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
        ]
        fecha_actual = datetime.now()
        fecha_str = f"TALCA, {fecha_actual.day} de {meses[fecha_actual.month - 1]} de {anio}"

        # Encabezado: logo + epigrafe
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        left_cell.width = Inches(3.2)
        right_cell.width = Inches(3.4)
        left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        logo_path = find('img/image_5e2339.png')
        if logo_path and os.path.exists(logo_path):
            p_logo = left_cell.paragraphs[0]
            run_logo = p_logo.add_run()
            run_logo.add_picture(logo_path, width=Inches(1.8))

        p = right_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_compact(p)
        add_paragraph_border(p, top=True, bottom=True)

        run = p.add_run(titulo)
        set_run_font(run)

        p = right_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(fecha_str)
        set_run_font(run)

        p = right_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"N° {numero_ru}")
        set_run_font(run, bold=True)

        doc.add_paragraph()

        add_section_title("VISTOS:")

        vistos = "; ".join([" ".join(v.strip().split()) for v in vistos_textos if v.strip()])
        if vistos:
            vistos += "."
        else:
             vistos = "[No se seleccionaron antecedentes normativos o vistos en el asistente]."

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_compact(p)

        run = p.add_run(vistos)
        set_run_font(run)

        add_section_title("CONSIDERANDO:")

        lista_considerandos = [c.strip() for c in considerandos_extra if c.strip()]
        for idx, cons_texto in enumerate(lista_considerandos):
            letra = chr(97 + idx)
            cons_limpio = " ".join(cons_texto.split())

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(3.2)
            p.paragraph_format.space_after = Pt(8)

            run = p.add_run(f"{letra}) {cons_limpio}")
            set_run_font(run)

        add_section_title("RESUELVO:")

        if resuelvo_texto:
            for linea in resuelvo_texto.splitlines():
                if linea.strip():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run = p.add_run(linea.strip())
                    set_run_font(run)
        else:
            p = doc.add_paragraph()
            run = p.add_run("[Acción resolutiva principal no ingresada].")
            set_run_font(run)

        doc.add_paragraph()
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ANÓTESE Y COMUNÍQUESE")
        set_run_font(run)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('"POR ORDEN DEL SR. RECTOR"')
        set_run_font(run)

        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("SWW/")
        set_run_font(run)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="RU_{numero_ru}-{anio}.docx"'
        return response